import docx
import MTK2
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX

doc = docx.Document('variant03.docx')

if __name__ == '__main__':
    text = "бог сделал людей, кольт сделал их равными"

    TextMTK2 = MTK2.MTK2_code(text)
    print(type(TextMTK2))

    LenParagraphs = []
    OpenText = ""
    for paragraph in doc.paragraphs:
        stroka = ""
        for run in paragraph.runs:
            for char in run.text:
                stroka += char
        OpenText += stroka
        LenParagraphs.append(len(stroka))
    print('stroka', stroka)

    print('OT', OpenText)
    print('length P', LenParagraphs)


    doc.paragraphs.clear()
    id_char = 0
    for id_paragraph in range(len(doc.paragraphs)):
        doc.paragraphs[id_paragraph].clear()
        for id_rans in range(LenParagraphs[id_paragraph]):
            run = doc.paragraphs[id_paragraph].add_run(OpenText[id_char])
            if id_char < len(TextMTK2):
                if TextMTK2[id_char] == '1':
                    run.font.color.rgb = RGBColor(1, 0, 1)
                else:
                    run.font.color.rgb = RGBColor(0, 0, 0)
            else:
                run.font.color.rgb = RGBColor(0, 0, 0)
            id_char += 1


    doc.save('variant03.docx')
